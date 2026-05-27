"""
paper2trace: research-decision trace extraction for papers and reports

用法：
  python paper2trace.py <文件路径>   # 支持 .docx / .pdf / .txt

配置：
  - 默认值在代码中保持安全通用默认
  - 私有配置写入仓库根目录 .env（可参考 .env.example）
  - 环境变量优先级高于 .env
"""
import argparse
import copy
import json
import os
import random
import re
import sys
import time
from pathlib import Path

import docx

from prompts import (
    DOMAIN_EXPERT_SYSTEM,
    EXTRACT_STRUCTURE,
    EXTRACT_RAW_DATA,
    EXTRACT_HYPOTHESIS_CHAIN,
    CRITICAL_ANALYSIS,
    GENERATE_SFT,
    GENERATE_DPO,
    GENERATE_REACT,
    SEMANTIC_QUALITY_CHECK,
)

PAPER2TRACE_DIR = Path(__file__).parent
REPO_ROOT = PAPER2TRACE_DIR.parent
DEFAULT_ENV_PATH = REPO_ROOT / ".env"

DEFAULT_CONFIG = {
    "llm": {
        "provider": "anthropic",
        "model": "",
        "temperature": 0.2,
        "default_max_tokens": 8192,
        "anthropic": {
            "api_key": "",
        },
        "openai": {
            "api_key": "",
            "base_url": "https://api.openai.com/v1",
        },
        "step_max_tokens": {
            "sft": 16384,
            "dpo": 16384,
        },
    },
    "paper2trace": {
        "output_dir": "paper2trace/output",
        "document_max_chars": 28000,
        "paper_id_max_length": 30,
        "random_seed": 0,
    },
    "semantic_quality": {
        "inferred_sample_size": 3,
        "sft_sample_size": 4,
        "dpo_sample_size": 3,
    },
    "quality_thresholds": {
        "min_rejected_nodes": 2,
        "min_inferred_nodes": 3,
        "min_sft_items": 15,
        "min_react_thoughts": 5,
        "min_dpo_chosen_chars": 300,
        "min_dpo_length_ratio": 0.45,
    },
}

CONFIG = copy.deepcopy(DEFAULT_CONFIG)
OUTPUT_DIR = REPO_ROOT / CONFIG["paper2trace"]["output_dir"]
_PROVIDER = CONFIG["llm"]["provider"]
MODEL = CONFIG["llm"]["model"]
_client = None


def _deep_merge(base: dict, override: dict) -> dict:
    merged = copy.deepcopy(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = _deep_merge(merged[key], value)
        elif value is not None:
            merged[key] = value
    return merged


def _load_env_file(path: Path) -> dict:
    if not path.exists():
        return {}
    values = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if not stripped or stripped.startswith("#") or "=" not in stripped:
            continue
        key, value = stripped.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        values[key] = value
    return values


def _env_int(env: dict, name: str):
    value = env.get(name)
    return int(value) if value not in (None, "") else None


def _env_float(env: dict, name: str):
    value = env.get(name)
    return float(value) if value not in (None, "") else None


def _env_overrides(env: dict) -> dict:
    return {
        "llm": {
            "provider": env.get("API_PROVIDER"),
            "model": env.get("MODEL"),
            "temperature": _env_float(env, "LLM_TEMPERATURE"),
            "default_max_tokens": _env_int(env, "LLM_MAX_TOKENS"),
            "anthropic": {
                "api_key": env.get("ANTHROPIC_API_KEY"),
            },
            "openai": {
                "api_key": env.get("OPENAI_API_KEY"),
                "base_url": env.get("OPENAI_BASE_URL"),
            },
            "step_max_tokens": {
                "sft": _env_int(env, "LLM_SFT_MAX_TOKENS"),
                "dpo": _env_int(env, "LLM_DPO_MAX_TOKENS"),
            },
        },
        "paper2trace": {
            "output_dir": env.get("PAPER2TRACE_OUTPUT_DIR"),
            "document_max_chars": _env_int(env, "PAPER2TRACE_DOCUMENT_MAX_CHARS"),
            "paper_id_max_length": _env_int(env, "PAPER2TRACE_PAPER_ID_MAX_LENGTH"),
            "random_seed": _env_int(env, "PAPER2TRACE_RANDOM_SEED"),
        },
        "semantic_quality": {
            "inferred_sample_size": _env_int(env, "SEMANTIC_INFERRED_SAMPLE_SIZE"),
            "sft_sample_size": _env_int(env, "SEMANTIC_SFT_SAMPLE_SIZE"),
            "dpo_sample_size": _env_int(env, "SEMANTIC_DPO_SAMPLE_SIZE"),
        },
        "quality_thresholds": {
            "min_rejected_nodes": _env_int(env, "QUALITY_MIN_REJECTED_NODES"),
            "min_inferred_nodes": _env_int(env, "QUALITY_MIN_INFERRED_NODES"),
            "min_sft_items": _env_int(env, "QUALITY_MIN_SFT_ITEMS"),
            "min_react_thoughts": _env_int(env, "QUALITY_MIN_REACT_THOUGHTS"),
            "min_dpo_chosen_chars": _env_int(env, "QUALITY_MIN_DPO_CHOSEN_CHARS"),
            "min_dpo_length_ratio": _env_float(env, "QUALITY_MIN_DPO_LENGTH_RATIO"),
        },
    }


def configure(env_path: str = None) -> None:
    global CONFIG, OUTPUT_DIR, _PROVIDER, MODEL, _client

    env_file_values = _load_env_file(Path(env_path) if env_path else DEFAULT_ENV_PATH)
    env_values = {**env_file_values, **os.environ}
    CONFIG = _deep_merge(DEFAULT_CONFIG, _env_overrides(env_values))

    _PROVIDER = CONFIG["llm"]["provider"].lower()
    MODEL = CONFIG["llm"].get("model") or (
        "claude-sonnet-4-6" if _PROVIDER == "anthropic" else "gpt-4o"
    )
    CONFIG["llm"]["model"] = MODEL
    output_dir = Path(CONFIG["paper2trace"]["output_dir"])
    OUTPUT_DIR = output_dir if output_dir.is_absolute() else REPO_ROOT / output_dir
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    random.seed(CONFIG["paper2trace"].get("random_seed", 0))
    _client = None


def _init_llm_client():
    global _client
    if _client is not None:
        return _client

    if _PROVIDER == "anthropic":
        api_key = CONFIG["llm"]["anthropic"].get("api_key")
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY 未配置；请写入 .env 或设置环境变量")
        import anthropic as _anthropic
        _client = _anthropic.Anthropic(api_key=api_key)

    elif _PROVIDER == "openai":
        api_key = CONFIG["llm"]["openai"].get("api_key")
        if not api_key:
            raise ValueError("OPENAI_API_KEY 未配置；请写入 .env 或设置环境变量")
        from openai import OpenAI as _OpenAI
        _client = _OpenAI(
            api_key=api_key,
            base_url=CONFIG["llm"]["openai"].get("base_url") or "https://api.openai.com/v1",
        )

    else:
        raise ValueError(f"不支持的 API_PROVIDER: {_PROVIDER}，可选：anthropic / openai")

    return _client


def _step_max_tokens(step_key: str) -> int:
    return CONFIG["llm"].get("step_max_tokens", {}).get(
        step_key,
        CONFIG["llm"].get("default_max_tokens", 8192),
    )


configure()


def paper_output_dir(paper_id: str) -> Path:
    """Return the output directory for a single paper run."""
    return OUTPUT_DIR / paper_id

# ─────────────────────────────────────────────────────────────────────────────
# 文档解析（支持 .docx / .pdf / .txt）
# ─────────────────────────────────────────────────────────────────────────────
def _parse_docx(filepath: str) -> str:
    doc = docx.Document(filepath)
    parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _parse_pdf(filepath: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        print("  [错误] 读取 PDF 需要安装 pdfplumber：pip install pdfplumber")
        sys.exit(1)
    parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return "\n".join(parts)


def _parse_txt(filepath: str) -> str:
    return Path(filepath).read_text(encoding="utf-8")


def parse_document(filepath: str) -> str:
    suffix = Path(filepath).suffix.lower()
    if suffix == ".docx":
        return _parse_docx(filepath)
    elif suffix == ".pdf":
        return _parse_pdf(filepath)
    elif suffix == ".txt":
        return _parse_txt(filepath)
    else:
        print(f"  [警告] 不支持的文件格式 {suffix}，尝试按纯文本读取")
        return _parse_txt(filepath)


def truncate(text: str, max_chars: int = None) -> str:
    """超长文档截断避免浪费 token；PDF 页数多时建议分批处理"""
    max_chars = max_chars or CONFIG["paper2trace"].get("document_max_chars", 28000)
    if len(text) <= max_chars:
        return text
    print(f"  [警告] 文档长度 {len(text)} 超过 {max_chars}，已截断")
    return text[:max_chars] + "\n...[文档已截断]"


# ─────────────────────────────────────────────────────────────────────────────
# LLM 调用（统一接口，屏蔽 Anthropic / OpenAI 差异）
# ─────────────────────────────────────────────────────────────────────────────
def call_llm(prompt: str, step_name: str, max_tokens: int = None, system: str = "") -> str:
    print(f"  → [{_PROVIDER}] {step_name}...", end="", flush=True)
    t0 = time.time()
    client = _init_llm_client()
    max_tokens = max_tokens or CONFIG["llm"].get("default_max_tokens", 8192)

    if _PROVIDER == "anthropic":
        kwargs = dict(model=MODEL, max_tokens=max_tokens,
                      messages=[{"role": "user", "content": prompt}])
        if system:
            kwargs["system"] = system
        msg = client.messages.create(**kwargs)
        content   = msg.content[0].text
        tok_in    = msg.usage.input_tokens
        tok_out   = msg.usage.output_tokens

    else:  # openai compatible
        messages = []
        if system:
            messages.append({"role": "system", "content": system})
        messages.append({"role": "user", "content": prompt})
        resp = client.chat.completions.create(
            model=MODEL,
            max_tokens=max_tokens,
            messages=messages,
            temperature=CONFIG["llm"].get("temperature", 0.2),
        )
        content   = resp.choices[0].message.content
        usage     = resp.usage
        tok_in    = getattr(usage, "prompt_tokens", "?")
        tok_out   = getattr(usage, "completion_tokens", "?")

    elapsed = time.time() - t0
    print(f" 完成 ({elapsed:.1f}s, in={tok_in}, out={tok_out})")
    return content


def parse_json_output(raw: str, step_name: str, paper_id: str = None) -> dict:
    """从 LLM 输出中提取 JSON，容忍 markdown 代码块包装"""
    # 去掉 ```json ... ``` 包装
    cleaned = re.sub(r"^```(?:json)?\s*", "", raw.strip())
    cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError as e:
        print(f"  [警告] {step_name} JSON 解析失败: {e}")
        print(f"  原始输出前 200 字符: {raw[:200]}")
        # 保存原始输出供调试
        debug_dir = paper_output_dir(paper_id) if paper_id else OUTPUT_DIR
        debug_dir.mkdir(parents=True, exist_ok=True)
        debug_path = debug_dir / f"debug_{step_name}.txt"
        debug_path.write_text(raw, encoding="utf-8")
        raise


def save(paper_id: str, step: str, data) -> Path:
    suffix = ".txt" if isinstance(data, str) else ".json"
    out_dir = paper_output_dir(paper_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{step}{suffix}"
    if isinstance(data, str):
        path.write_text(data, encoding="utf-8")
    else:
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"  ✓ 保存 → {paper_id}/{path.name}")
    return path


def load(paper_id: str, step: str):
    for suffix in [".json", ".txt"]:
        path = paper_output_dir(paper_id) / f"{step}{suffix}"
        if path.exists():
            if suffix == ".json":
                return json.loads(path.read_text(encoding="utf-8"))
            else:
                return path.read_text(encoding="utf-8")
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Pipeline 各步骤
# ─────────────────────────────────────────────────────────────────────────────
def step_structure(paper_id: str, text: str) -> dict:
    print("\n[Step 1] 结构提取")
    cached = load(paper_id, "1_structure")
    if cached:
        print("  (使用缓存)")
        return cached
    raw = call_llm(EXTRACT_STRUCTURE.format(text=text), "结构提取")
    result = parse_json_output(raw, "结构提取", paper_id=paper_id)
    save(paper_id, "1_structure", result)
    return result


def step_raw_data(paper_id: str, text: str, system_prompt: str = "") -> dict:
    print("\n[Step 1.5] 精细数据提取（知识基底）")
    cached = load(paper_id, "1b_raw_data")
    if cached:
        print("  (使用缓存)")
        return cached
    raw = call_llm(EXTRACT_RAW_DATA.format(text=text), "精细数据提取", system=system_prompt)
    result = parse_json_output(raw, "精细数据提取", paper_id=paper_id)
    save(paper_id, "1b_raw_data", result)
    _print_raw_data_summary(result)
    return result


def _print_raw_data_summary(data: dict):
    """打印精细数据提取的关键统计，便于快速判断质量"""
    tables    = data.get("tables", [])
    claims    = data.get("quantitative_claims", [])
    params    = data.get("conditions_and_parameters", [])
    rejected  = data.get("tried_and_rejected", [])
    steps     = data.get("procedural_steps", [])
    gaps      = data.get("information_gaps", [])

    total_rows = sum(len(t.get("rows", [])) for t in tables)
    explicit_c = sum(1 for x in claims + params if x.get("source") == "explicit")
    inferred_c = sum(1 for x in claims + params if x.get("source") == "inferred")

    print(f"  数据表:     {len(tables)} 张，共 {total_rows} 行")
    print(f"  定量断言:   {len(claims)} 条（explicit={explicit_c}, inferred={inferred_c}）")
    print(f"  条件参数:   {len(params)} 条")
    print(f"  尝试后淘汰: {len(rejected)} 条")
    print(f"  操作序列:   {len(steps)} 条")
    print(f"  信息缺口:   {len(gaps)} 条")

    if gaps:
        print("  ── 信息缺口:")
        for g in gaps:
            print(f"     · {g.get('gap', '')}  [{g.get('likely_reason', '')}]")


def step_hypothesis_chain(paper_id: str, text: str, structure: dict, raw_data: dict,
                          system_prompt: str = "") -> dict:
    print("\n[Step 2] 假设链提取")
    cached = load(paper_id, "2_hypothesis_chain")
    if cached:
        print("  (使用缓存)")
        return cached
    prompt = EXTRACT_HYPOTHESIS_CHAIN.format(
        text=text,
        structure=json.dumps(structure, ensure_ascii=False, indent=2),
        raw_data=json.dumps(raw_data, ensure_ascii=False, indent=2),
    )
    raw = call_llm(prompt, "假设链提取", system=system_prompt)
    result = parse_json_output(raw, "假设链提取", paper_id=paper_id)
    save(paper_id, "2_hypothesis_chain", result)
    return result


def step_critical_analysis(paper_id: str, text: str, structure: dict, hypothesis: dict,
                            system_prompt: str = "") -> dict:
    print("\n[Step 3] 批判性分析")
    cached = load(paper_id, "3_critical_analysis")
    if cached:
        print("  (使用缓存)")
        return cached
    domain = structure.get("domain", "科研")
    prompt = CRITICAL_ANALYSIS.format(
        domain=domain,
        text=text,
        hypothesis_chain=json.dumps(hypothesis, ensure_ascii=False, indent=2),
    )
    raw = call_llm(prompt, "批判性分析", system=system_prompt)
    result = parse_json_output(raw, "批判性分析", paper_id=paper_id)
    save(paper_id, "3_critical_analysis", result)
    return result


def step_sft(paper_id: str, structure: dict, raw_data: dict,
             hypothesis: dict, critical: dict, system_prompt: str = "") -> dict:
    print("\n[Step 4a] 生成 SFT 数据")
    cached = load(paper_id, "4a_sft")
    if cached:
        print("  (使用缓存)")
        return cached
    prompt = GENERATE_SFT.format(
        structure=json.dumps(structure, ensure_ascii=False, indent=2),
        raw_data=json.dumps(raw_data, ensure_ascii=False, indent=2),
        hypothesis_chain=json.dumps(hypothesis, ensure_ascii=False, indent=2),
        critical_analysis=json.dumps(critical, ensure_ascii=False, indent=2),
    )
    raw = call_llm(prompt, "SFT生成", max_tokens=_step_max_tokens("sft"), system=system_prompt)
    result = parse_json_output(raw, "SFT生成", paper_id=paper_id)
    save(paper_id, "4a_sft", result)
    return result


def step_dpo(paper_id: str, hypothesis: dict, critical: dict, system_prompt: str = "") -> dict:
    print("\n[Step 4b] 生成 DPO 数据")
    cached = load(paper_id, "4b_dpo")
    if cached:
        print("  (使用缓存)")
        return cached
    prompt = GENERATE_DPO.format(
        hypothesis_chain=json.dumps(hypothesis, ensure_ascii=False, indent=2),
        critical_analysis=json.dumps(critical, ensure_ascii=False, indent=2),
    )
    raw = call_llm(prompt, "DPO生成", max_tokens=_step_max_tokens("dpo"), system=system_prompt)
    result = parse_json_output(raw, "DPO生成", paper_id=paper_id)
    save(paper_id, "4b_dpo", result)
    return result


def step_react(paper_id: str, structure: dict, hypothesis: dict, system_prompt: str = "") -> str:
    print("\n[Step 4c] 生成科研过程推理 CoT")
    cached = load(paper_id, "4c_react")
    if cached:
        print("  (使用缓存)")
        return cached
    prompt = GENERATE_REACT.format(
        structure=json.dumps(structure, ensure_ascii=False, indent=2),
        hypothesis_chain=json.dumps(hypothesis, ensure_ascii=False, indent=2),
    )
    raw = call_llm(prompt, "过程推理CoT生成", system=system_prompt)
    save(paper_id, "4c_react", raw)
    return raw


# ─────────────────────────────────────────────────────────────────────────────
# 语义质量检查（LLM 抽样审核）
# ─────────────────────────────────────────────────────────────────────────────
def step_semantic_quality(paper_id: str, hypothesis: dict, sft: dict, dpo: dict) -> dict:
    print("\n[Step 5b] 语义质量检查")
    cached = load(paper_id, "5b_semantic_quality")
    if cached:
        print("  (使用缓存)")
        return cached

    chain = hypothesis.get("hypothesis_chain", [])
    sft_items = sft.get("sft_data", [])
    dpo_items = dpo.get("dpo_data", [])

    inferred = [h for h in chain if h.get("information_source") == "inferred"]
    sample_config = CONFIG["semantic_quality"]
    inferred_sample = random.sample(
        inferred,
        min(sample_config["inferred_sample_size"], len(inferred)),
    )
    sft_sample = random.sample(sft_items, min(sample_config["sft_sample_size"], len(sft_items)))
    dpo_sample = random.sample(dpo_items, min(sample_config["dpo_sample_size"], len(dpo_items)))

    prompt = SEMANTIC_QUALITY_CHECK.format(
        inferred_nodes_sample=json.dumps(inferred_sample, ensure_ascii=False, indent=2),
        sft_sample=json.dumps(sft_sample, ensure_ascii=False, indent=2),
        dpo_sample=json.dumps(dpo_sample, ensure_ascii=False, indent=2),
    )
    raw = call_llm(prompt, "语义质检")
    result = parse_json_output(raw, "语义质检", paper_id=paper_id)
    save(paper_id, "5b_semantic_quality", result)

    overall = result.get("overall", {})
    score = overall.get("score", "?")
    passed = overall.get("pass", False)
    print(f"  综合得分: {score}/5  {'✓ 通过' if passed else '✗ 未通过'}")
    print(f"  建议: {overall.get('recommendation', '')}")
    return result


# ─────────────────────────────────────────────────────────────────────────────
# 规则质量检查
# ─────────────────────────────────────────────────────────────────────────────
def quality_check(paper_id: str, hypothesis: dict, sft: dict, dpo: dict, react: str):
    print("\n[质量检查]")
    chain = hypothesis.get("hypothesis_chain", [])
    sft_items  = sft.get("sft_data", [])
    dpo_items  = dpo.get("dpo_data", [])

    rejected = [h for h in chain if h.get("decision") == "rejected"]
    accepted = [h for h in chain if h.get("decision") == "accepted"]
    inferred = [h for h in chain if h.get("information_source") == "inferred"]

    report = {
        "paper_id": paper_id,
        "hypothesis_chain": {
            "total_nodes": len(chain),
            "rejected": len(rejected),
            "accepted": len(accepted),
            "inferred_nodes": len(inferred),
        },
        "sft": {
            "total": len(sft_items),
            "types": {},
        },
        "dpo": {
            "total": len(dpo_items),
        },
        "react": {
            "length_chars": len(react),
            "observation_count": react.count("[观察]"),
            "thought_count": react.count("[思考]"),
            "action_count": react.count("[行动]"),
        },
        "warnings": [],
    }

    for item in sft_items:
        t = item.get("type", "unknown")
        report["sft"]["types"][t] = report["sft"]["types"].get(t, 0) + 1

    # 基础质检规则
    thresholds = CONFIG["quality_thresholds"]
    if len(rejected) < thresholds["min_rejected_nodes"]:
        report["warnings"].append("假设链中被淘汰节点过少，过程数据可能不够丰富")
    if len(inferred) < thresholds["min_inferred_nodes"]:
        report["warnings"].append(
            f"推断型节点仅{len(inferred)}个（期望≥{thresholds['min_inferred_nodes']}），隐性过程数据挖掘不足"
        )
    if len(sft_items) < thresholds["min_sft_items"]:
        report["warnings"].append(f"SFT数据仅{len(sft_items)}条，偏少")
    if report["react"]["thought_count"] < thresholds["min_react_thoughts"]:
        report["warnings"].append("过程推理CoT思考步骤过少，推理深度不够")
    for item in dpo_items:
        chosen_len = len(item.get("chosen", {}).get("response", ""))
        rejected_len = len(item.get("rejected", {}).get("response", ""))
        ratio = min(chosen_len, rejected_len) / max(chosen_len, rejected_len) if max(chosen_len, rejected_len) > 0 else 0
        if chosen_len < thresholds["min_dpo_chosen_chars"]:
            report["warnings"].append(
                f"DPO {item.get('id')} chosen 仅{chosen_len}字（期望≥{thresholds['min_dpo_chosen_chars']}）"
            )
        if ratio < thresholds["min_dpo_length_ratio"]:
            report["warnings"].append(f"DPO {item.get('id')} chosen/rejected 长度差距过大（比例{ratio:.2f}），模型可能通过长度判断好坏")

    print(f"  假设链节点: {len(chain)} 个（淘汰 {len(rejected)}，推断 {len(inferred)}）")
    print(f"  SFT: {len(sft_items)} 条  |  DPO: {len(dpo_items)} 对  |  过程推理CoT: {report['react']['observation_count']} 轮")
    if report["warnings"]:
        print("  ⚠ 规则警告:")
        for w in report["warnings"]:
            print(f"    - {w}")
    else:
        print("  ✓ 规则检查通过")

    return report


def merge_quality_report(paper_id: str, rule_report: dict, semantic_report: dict):
    """将规则质检和语义质检合并成一份最终报告"""
    merged = dict(rule_report)
    merged["semantic_quality"] = semantic_report
    merged["overall_pass"] = (
        len(rule_report.get("warnings", [])) == 0
        and semantic_report.get("overall", {}).get("pass", False)
    )
    save(paper_id, "5_quality_report", merged)
    return merged


# ─────────────────────────────────────────────────────────────────────────────
# 主流程
# ─────────────────────────────────────────────────────────────────────────────
def run(filepath: str, steps: str = "all", paper_id: str = None):
    path = Path(filepath)
    if not path.exists():
        print(f"文件不存在: {filepath}")
        sys.exit(1)

    paper_id_max_length = CONFIG["paper2trace"].get("paper_id_max_length", 30)
    paper_id = paper_id or path.stem.replace(" ", "_")[:paper_id_max_length]
    print(f"\n{'='*60}")
    print(f"paper2trace 启动")
    print(f"  文件: {path.name}")
    print(f"  paper_id: {paper_id}")
    print(f"{'='*60}")

    # 解析文档
    print("\n[解析文档]")
    raw_text = parse_document(filepath)
    text = truncate(raw_text)
    print(f"  文档字数: {len(raw_text)}，处理字数: {len(text)}")

    # 运行各步骤
    structure  = step_structure(paper_id, text)

    # Step 1 完成后，用提取到的领域构建专家人设 system prompt，注入后续所有步骤
    domain = structure.get("domain", "科研")
    sys_prompt = DOMAIN_EXPERT_SYSTEM.format(domain=domain)
    print(f"\n  [领域专家身份] {domain}")

    raw_data   = step_raw_data(paper_id, text, system_prompt=sys_prompt)
    hypothesis = step_hypothesis_chain(paper_id, text, structure, raw_data, system_prompt=sys_prompt)
    critical   = step_critical_analysis(paper_id, text, structure, hypothesis, system_prompt=sys_prompt)
    sft        = step_sft(paper_id, structure, raw_data, hypothesis, critical, system_prompt=sys_prompt)
    dpo        = step_dpo(paper_id, hypothesis, critical, system_prompt=sys_prompt)
    react      = step_react(paper_id, structure, hypothesis, system_prompt=sys_prompt)

    # 质量检查（规则 + 语义）
    print("\n[Step 5a] 规则质量检查")
    rule_report = quality_check(paper_id, hypothesis, sft, dpo, react)
    semantic_report = step_semantic_quality(paper_id, hypothesis, sft, dpo)
    final_report = merge_quality_report(paper_id, rule_report, semantic_report)

    passed = final_report.get("overall_pass", False)
    print(f"\n  最终质检: {'✓ 通过' if passed else '✗ 未通过（见 5_quality_report.json）'}")

    print(f"\n{'='*60}")
    print(f"paper2trace 完成！输出目录: {paper_output_dir(paper_id)}")
    print(f"{'='*60}\n")

    return {
        "structure": structure,
        "hypothesis_chain": hypothesis,
        "critical_analysis": critical,
        "sft": sft,
        "dpo": dpo,
        "react": react,
        "quality_report": final_report,
    }


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="paper2trace 科研决策轨迹抽取器")
    parser.add_argument("filepath", help="报告文件路径（支持 .docx / .pdf / .txt）")
    parser.add_argument("--paper-id", help="自定义 paper_id（默认用文件名）")
    parser.add_argument("--env-file", help="自定义 .env 文件路径（默认读取仓库根目录 .env）")
    args = parser.parse_args()
    if args.env_file:
        configure(env_path=args.env_file)
    run(args.filepath, paper_id=args.paper_id)
